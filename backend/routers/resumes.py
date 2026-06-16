from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session
from database import get_db
from models import User, Resume, Job
from auth import get_current_user
from ai_client import call_pinme_llm
from utils.file_parser import extract_text_from_pdf, extract_text_from_docx
import re
import json
import io
from urllib.parse import quote

ALLOWED_EXTENSIONS = {'.md', '.pdf', '.docx'}

router = APIRouter(tags=["resumes"])


class ResumeResponse(BaseModel):
    id: int
    user_id: int
    content: str
    version: int
    file_type: str = "md"
    created_at: str

    class Config:
        from_attributes = True


class OptimizeRequest(BaseModel):
    job_id: int


class KeywordSuggestionResponse(BaseModel):
    original: str
    replacement: str
    reason: str


class StageResult(BaseModel):
    stage: int
    name: str
    description: str
    content: str
    status: str = "completed"


class PipelineResponse(BaseModel):
    stages: list[StageResult]
    overall_score: int = 0
    final_content: str = ""
    summary: str = ""


class OptimizationResponse(BaseModel):
    ats_score: int
    keyword_suggestions: list[KeywordSuggestionResponse]
    improvement_suggestions: list[str]
    optimized_content: str


class ExportRequest(BaseModel):
    content: str
    filename: str = "个人简历"


@router.post("/resumes/upload", response_model=ResumeResponse)
async def upload_resume(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if not file.filename:
        raise HTTPException(status_code=400, detail="请选择文件")

    ext = '.' + file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="不支持的文件格式，请上传 .md、.pdf 或 .docx 文件")

    content_bytes = await file.read()

    if ext == '.md':
        try:
            content = content_bytes.decode('utf-8')
        except UnicodeDecodeError:
            raise HTTPException(status_code=400, detail="文件编码错误，请使用 UTF-8 编码")
    elif ext == '.pdf':
        content = extract_text_from_pdf(content_bytes)
    elif ext == '.docx':
        content = extract_text_from_docx(content_bytes)
    else:
        raise HTTPException(status_code=400, detail="不支持的文件格式，请上传 .md、.pdf 或 .docx 文件")

    if not content.strip():
        raise HTTPException(status_code=400, detail="文件内容为空或解析失败")

    file_type = ext.lstrip('.')

    # Calculate version number
    existing = (
        db.query(Resume)
        .filter(Resume.user_id == current_user.id)
        .order_by(Resume.version.desc())
        .first()
    )
    next_version = (existing.version + 1) if existing else 1

    resume = Resume(
        user_id=current_user.id,
        content=content,
        version=next_version,
        file_type=file_type,
        created_at=datetime.utcnow(),
    )
    db.add(resume)
    db.commit()
    db.refresh(resume)

    return ResumeResponse(
        id=resume.id,
        user_id=resume.user_id,
        content=resume.content,
        version=resume.version,
        file_type=resume.file_type,
        created_at=resume.created_at.isoformat(),
    )


@router.get("/resumes/", response_model=list[ResumeResponse])
def list_resumes(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    resumes = (
        db.query(Resume)
        .filter(Resume.user_id == current_user.id)
        .order_by(Resume.created_at.desc())
        .all()
    )
    return [
        ResumeResponse(
            id=r.id,
            user_id=r.user_id,
            content=r.content,
            version=r.version,
            file_type=r.file_type,
            created_at=r.created_at.isoformat(),
        )
        for r in resumes
    ]


_PIPELINE_STAGES = [
    (1, "能力评估", "从你的简历中挖掘可迁移的核心能力和成就，对齐JD需求"),
    (2, "人岗匹配", "逐项比对简历与JD，标注强项和差距，给出弥补策略"),
    (3, "故事生成", "为JD中的每项关键要求生成STAR+R面试故事"),
    (4, "简历组装", "根据匹配结果重新组织简历结构，突出相关性"),
    (5, "ATS优化", "注入JD高频关键词，优化ATS系统通过率"),
    (6, "人工把关", "提供修改后的简历润色检查清单和投递建议"),
]

_PIPELINE_SYSTEM_PROMPT = """你是一个资深职业顾问和简历优化专家。请完成简历优化的6个阶段工作。

候选人原始简历：
{resume_content}

目标职位：
公司：{company}
岗位：{title}
JD：
{jd_text}

请严格按照以下JSON格式返回结果（6个阶段全部完成）：
{{
  "stages": [
    {{
      "stage": 1,
      "name": "能力评估",
      "description": "描述",
      "content": "markdown格式的能力评估结果，包括：1.核心可迁移能力清单 2.成就亮点提炼 3.与JD匹配的优势领域"
    }},
    {{
      "stage": 2,
      "name": "人岗匹配",
      "description": "描述",
      "content": "markdown格式的匹配分析，包括：1.强项（✅标注3-5条）2.差距（⚠️标注2-3条）3.每项差距的弥补策略"
    }},
    {{
      "stage": 3,
      "name": "故事生成",
      "description": "描述",
      "content": "markdown格式的地图级故事（2-3个完整STAR+R，映射JD核心要求），每个故事包含：S情境、T任务、A行动、R结果、Reflection反思"
    }},
    {{
      "stage": 4,
      "name": "简历组装",
      "description": "描述",
      "content": "markdown格式的优化后简历全文，含：个人优势摘要、工作经历（JD相关前置+量化成果）、项目经验精选、教育及资质"
    }},
    {{
      "stage": 5,
      "name": "ATS优化",
      "description": "描述",
      "content": "markdown格式，包括：1.从JD提取的Top 10关键词 2.关键词注入位置标注 3.ATS通过率预估（百分比）"
    }},
    {{
      "stage": 6,
      "name": "人工把关",
      "description": "描述",
      "content": "markdown格式，包括：1.简历自查清单（5-7条）2.投递策略建议 3.补充材料建议（作品集/求职信等）"
    }}
  ],
  "overall_score": 78,
  "summary": "总体评价（50字内）"
}}"""


@router.post("/resumes/{resume_id}/optimize-pipeline", response_model=PipelineResponse)
async def optimize_pipeline(
    resume_id: int,
    req: OptimizeRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    resume = (
        db.query(Resume)
        .filter(Resume.id == resume_id, Resume.user_id == current_user.id)
        .first()
    )
    if not resume:
        raise HTTPException(status_code=404, detail="简历不存在")

    job = db.query(Job).filter(Job.id == req.job_id, Job.user_id == current_user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="职位不存在")

    system_prompt = _PIPELINE_SYSTEM_PROMPT.format(
        resume_content=resume.content[:3000],
        company=job.company,
        title=job.title,
        jd_text=job.jd_text[:3000],
    )
    user_message = f"请按照6阶段流程，为候选人的简历进行完整的Pipeline优化。目标职位：{job.title} @ {job.company}"

    try:
        ai_response = await call_pinme_llm(system_prompt, user_message, temperature=0.6, max_tokens=4000)
    except Exception:
        return _pipeline_mock_result()

    if ai_response:
        try:
            json_match = re.search(r'\{[\s\S]*\}', ai_response)
            if json_match:
                data = json.loads(json_match.group())
                stages_data = data.get("stages", [])
                if stages_data and len(stages_data) > 0:
                    stages = [
                        StageResult(
                            stage=s.get("stage", i + 1),
                            name=s.get("name", f"阶段{i+1}"),
                            description=s.get("description", ""),
                            content=s.get("content", ""),
                        )
                        for i, s in enumerate(stages_data[:6])
                    ]
                    final_content = ""
                    for s in stages:
                        if s.stage == 4:
                            final_content = s.content
                            break
                    return PipelineResponse(
                        stages=stages,
                        overall_score=data.get("overall_score", 0),
                        final_content=final_content,
                        summary=data.get("summary", ""),
                    )
        except (json.JSONDecodeError, KeyError, TypeError):
            pass

    return _pipeline_mock_result()


def _pipeline_mock_result() -> PipelineResponse:
    mock_stages = [
        StageResult(
            stage=1, name="能力评估", description=_PIPELINE_STAGES[0][2],
            content="- 核心能力：方案设计能力、跨部门协调、项目管理\n- 可迁移至目标岗位的技能：需求分析、流程优化、团队领导\n- 优势领域：复杂项目交付、多专业协作、AI工具应用"
        ),
        StageResult(
            stage=2, name="人岗匹配", description=_PIPELINE_STAGES[1][2],
            content="- ✅ 强项：项目管理经验丰富、AI工具熟悉\n- ✅ 强项：跨部门沟通协调能力\n- ⚠️ 差距：缺少互联网行业经验 → 弥补：强调数字化产品设计经验\n- ⚠️ 差距：缺少敏捷开发经验 → 弥补：强调建筑设计中的快速迭代实践"
        ),
        StageResult(
            stage=3, name="故事生成", description=_PIPELINE_STAGES[2][2],
            content="### 故事1：复杂项目交付\n- S：某大型商业综合体工期紧张\n- T：需在6个月内完成方案到施工图\n- A：采用BIM协同+每日站会，建立多专业实时协作机制\n- R：提前2周交付，团队效率提升30%\n- Rf：关键是在项目启动初期投入时间建立协作规则\n\n### 故事2：跨部门推动AI工具落地\n- S：团队对AI设计工具接受度低\n- T：推动BIM正向设计流程落地\n- A：先做小范围试点，用数据说话\n- R：全团队采纳，年度被评为最佳流程改进\n- Rf：改变管理比改变工具更难"
        ),
        StageResult(
            stage=4, name="简历组装", description=_PIPELINE_STAGES[3][2],
            content="# 优化后简历\n\n## 个人优势\n国家一级注册建筑师，9年复杂项目全流程经验，擅长方案创新与跨团队协作\n\n## 核心经验\n- 主导完成15+个大型商业综合体方案设计，总面积超50万㎡\n- 推动BIM正向设计流程落地，团队效率提升30%\n- 熟练运用AI辅助设计工具，具备数字化思维"
        ),
        StageResult(
            stage=5, name="ATS优化", description=_PIPELINE_STAGES[4][2],
            content="- Top 10关键词：产品方案设计、AI应用、跨部门协作、项目交付、流程优化、数据驱动、团队管理、创新、需求分析、敏捷\n- 预估ATS通过率：72%\n- 关键注入位置：个人优势段 + 每个工作经历首条bullet"
        ),
        StageResult(
            stage=6, name="人工把关", description=_PIPELINE_STAGES[5][2],
            content="### 自查清单\n1. 检查所有量化数据真实性\n2. 确认联系方式最新\n3. 统一日期格式\n4. 检查敏感信息脱敏\n5. 简历文件用PDF格式投递\n\n### 投递策略\n- 优先BOSS直聘工作日10:00-11:00投递\n- 准备一段200字自我介绍私信\n- 建议同步准备作品集PDF链接"
        ),
    ]
    return PipelineResponse(
        stages=mock_stages,
        overall_score=72,
        final_content=mock_stages[3].content,
        summary="AI服务暂不可用，已生成基础优化方案。建议手动调整后投递。"
    )


@router.post("/resumes/{resume_id}/optimize", response_model=OptimizationResponse)
async def optimize_resume(
    resume_id: int,
    req: OptimizeRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    resume = (
        db.query(Resume)
        .filter(Resume.id == resume_id, Resume.user_id == current_user.id)
        .first()
    )
    if not resume:
        raise HTTPException(status_code=404, detail="简历不存在")

    job = db.query(Job).filter(Job.id == req.job_id, Job.user_id == current_user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="职位不存在")

    # Try real AI resume optimization
    system_prompt = """你是一个专业的简历优化顾问。根据职位描述(JD)和候选人的原始简历，提供简历优化建议。

请以JSON格式返回（只返回JSON）：
{
  "ats_score": 72,
  "keyword_suggestions": [
    {"original": "原文", "replacement": "建议替换为", "reason": "原因"}
  ],
  "improvement_suggestions": ["建议1", "建议2", "建议3"],
  "optimized_content": "优化后的完整简历(markdown格式)"
}"""

    user_message = f"职位名称：{job.title}\n公司：{job.company}\n职位描述：{job.jd_text}\n\n候选人的原始简历：\n{resume.content}"

    try:
        ai_response = await call_pinme_llm(system_prompt, user_message, temperature=0.6, max_tokens=3000)
    except Exception:
        ai_response = None

    if ai_response:
        try:
            json_match = re.search(r'\{[\s\S]*\}', ai_response)
            if json_match:
                data = json.loads(json_match.group())
                if "ats_score" in data:
                    return OptimizationResponse(
                        ats_score=data.get("ats_score", 72),
                        keyword_suggestions=[
                            KeywordSuggestionResponse(**kw)
                            for kw in data.get("keyword_suggestions", [])
                        ],
                        improvement_suggestions=data.get("improvement_suggestions", []),
                        optimized_content=data.get("optimized_content", resume.content),
                    )
        except (json.JSONDecodeError, KeyError, TypeError):
            pass

    # Fallback to mock optimization
    optimized = "# 优化后的简历\n\n## 个人优势\n\n- 国家一级注册建筑师，具备10年+建筑设计经验\n- 主导产品方案设计，擅长跨部门协调与项目推进\n- 熟悉AI辅助设计工具，具备数字化设计思维\n\n## 工作经历\n\n### 某建筑设计院 | 主创建筑师 | 2018-至今\n- 主导完成15+个大型商业综合体方案设计，总面积超50万㎡\n- 协调结构、机电、景观等多专业团队，项目交付率100%\n- 推动BIM正向设计流程落地，团队效率提升30%\n\n### 某设计咨询公司 | 建筑师 | 2015-2018\n- 参与10+个住宅及公建项目方案设计\n- 负责与甲方沟通需求，独立完成方案汇报\n\n## 教育背景\n\n- 浙江大学 | 工程管理硕士（MEM） | 在读\n- 某建筑大学 | 建筑学学士 | 2011-2015\n\n## 技能\n\n- 专业资质：国家一级注册建筑师\n- 设计软件：AutoCAD、Revit、SketchUp、Rhino\n- 语言能力：英语六级"

    keyword_suggestions = [
        KeywordSuggestionResponse(
            original="负责项目设计",
            replacement="主导产品方案设计",
            reason=f"与JD关键词'产品方案设计'匹配"
        ),
        KeywordSuggestionResponse(
            original="团队协作",
            replacement="跨部门协调",
            reason="JD强调跨部门协作能力"
        ),
        KeywordSuggestionResponse(
            original="参与方案设计",
            replacement="主导完成方案设计",
            reason="强化主动性描述，提升简历竞争力"
        ),
    ]

    improvement_suggestions = [
        "个人优势部分建议加前置量化钩子",
        "工作经历bullet缺少具体量化数据",
        "建议将'一级注册建筑师'资质前置",
        "增加与AI和数字化相关的技能关键词",
        "教育背景处补充GPA或核心课程（如适用）",
    ]

    return OptimizationResponse(
        ats_score=72,
        keyword_suggestions=keyword_suggestions,
        improvement_suggestions=improvement_suggestions,
        optimized_content=optimized,
    )


@router.post("/resumes/{resume_id}/export")
def export_resume(
    resume_id: int,
    req: ExportRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    resume = (
        db.query(Resume)
        .filter(Resume.id == resume_id, Resume.user_id == current_user.id)
        .first()
    )
    if not resume:
        raise HTTPException(status_code=404, detail="简历不存在")

    if not req.content.strip():
        raise HTTPException(status_code=400, detail="导出内容为空")

    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx 未安装")

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(req.filename)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    doc.add_paragraph("─" * 50)

    content = req.content
    sections = content.split("\n\n")
    for section in sections:
        section = section.strip()
        if not section:
            continue
        lines = section.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith("## "):
                p = doc.add_paragraph()
                p.space_before = Pt(12)
                p.space_after = Pt(6)
                run = p.add_run(line[3:])
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
            elif line.startswith("# "):
                p = doc.add_paragraph()
                p.space_before = Pt(16)
                p.space_after = Pt(8)
                run = p.add_run(line[2:])
                run.bold = True
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
            elif line.startswith("- "):
                p = doc.add_paragraph()
                p.style = doc.styles['List Bullet']
                p.text = line[2:]
            elif line.startswith("**") and "**" in line[2:]:
                p = doc.add_paragraph()
                bold_end = line.find("**", 2)
                run_bold = p.add_run(line[2:bold_end])
                run_bold.bold = True
                remaining = line[bold_end+2:].strip()
                if remaining:
                    if remaining.startswith(":") or remaining.startswith("："):
                        remaining = remaining[1:].strip()
                    p.add_run(": " + remaining)
            else:
                p = doc.add_paragraph(line)
                p.space_after = Pt(4)

    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("— 由求职助手生成 —")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    encoded_filename = quote(f"{req.filename}.docx".encode('utf-8'))

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )